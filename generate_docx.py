from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_bold_text(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(value)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()

# ---- TITLE ----
title = doc.add_heading('BaseCamp Digital (BCD) Website — Backend Requirements & Overview', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_bold_text(doc, 'Document Date: ', 'March 31, 2026')
add_bold_text(doc, 'Prepared By: ', 'BCD Product Team')
add_bold_text(doc, 'For: ', 'Tech / Infrastructure Team')

note = doc.add_paragraph()
note_run = note.add_run('📌 Note: Vimeo video setup, email (SMTP) configuration, and WordPress data migration are already completed. Those items are excluded from this document.')
note_run.italic = True

doc.add_paragraph()

# ---- PURPOSE ----
add_heading(doc, 'Purpose of This Document', 1)
doc.add_paragraph('This document has two parts:')
p = doc.add_paragraph(style='List Number')
p.add_run('Requirements FROM the Tech Team').bold = True
p.add_run(' — What we need the tech team to set up, configure, or hand over so we can build and operate the website.')
p = doc.add_paragraph(style='List Number')
p.add_run('What We Are Building on the Backend').bold = True
p.add_run(' — A plain-language overview of the backend systems we are implementing, so the tech team understands what they are supporting.')

# ---- PART 1 ----
add_heading(doc, 'Part 1: Requirements FROM the Tech Team', 1)
doc.add_paragraph(
    'The following items are blockers or dependencies that require action from the tech/infrastructure team. '
    'Please provide all credentials, configurations, and access in a secure manner (e.g., via a shared password manager or encrypted file).'
)

# Section 1
add_heading(doc, '1. Database — MySQL on SiteGround', 2)
doc.add_paragraph('The database is hosted on SiteGround using the team\'s own server. The following details are needed from the tech team:')
add_table(doc,
    ['#', 'Requirement', 'Details'],
    [
        ('1.1', 'MySQL connection credentials', 'Host, port, database name, username, and password. These go into the DATABASE_URL environment variable used by our Prisma ORM.'),
        ('1.2', 'Remote access / whitelist', 'Ensure the production server\'s IP is whitelisted in SiteGround to allow our application to connect to the MySQL database.'),
        ('1.3', 'Initial DB migration run', 'Run our Prisma migration scripts on the SiteGround MySQL database to create all required tables. We will provide the SQL/migration scripts.'),
        ('1.4', 'Seed the first admin user', 'After DB setup, manually set is_admin = true in the users table for the master admin account. No admin sign-up flow exists by design.'),
    ]
)

# Section 2
add_heading(doc, '2. Payment Gateway — Razorpay', 2)
add_table(doc,
    ['#', 'Requirement', 'Details'],
    [
        ('2.1', 'Razorpay Key ID', 'RAZORPAY_KEY_ID — public key used on the frontend checkout.'),
        ('2.2', 'Razorpay Key Secret', 'RAZORPAY_KEY_SECRET — secret key used server-side to create and verify payment orders. Must be kept confidential.'),
        ('2.3', 'Razorpay account activation', 'Ensure the Razorpay account is activated for live payments in INR. Test mode is fine for staging.'),
        ('2.4', 'Webhook configuration', 'Configure the Razorpay webhook to point to https://[domain]/api/payment-webhook. Provide the webhook secret as RAZORPAY_WEBHOOK_SECRET.'),
    ]
)

# Section 3
add_heading(doc, '3. Hosting & Deployment (Own Server)', 2)
doc.add_paragraph('The application is deployed on the team\'s own server. The tech team manages the infrastructure.')
add_table(doc,
    ['#', 'Requirement', 'Details'],
    [
        ('3.1', 'Server access / deployment pipeline', 'Provide SSH access or deployment credentials so we can push the Next.js build to the server. Alternatively, set up a CI/CD pipeline (e.g., GitHub Actions) to auto-deploy on push to main.'),
        ('3.2', 'Node.js environment', 'Ensure Node.js (v20+) and npm are installed on the server to run the Next.js application.'),
        ('3.3', 'Process manager', 'Set up a process manager (e.g., PM2) to keep the Next.js app running and restart it automatically on server reboot.'),
        ('3.4', 'Reverse proxy (Nginx/Apache)', 'Configure a reverse proxy (preferably Nginx) to forward traffic from port 80/443 to the Next.js app port (default: 3000).'),
        ('3.5', 'SSL certificate', 'Install an SSL certificate (e.g., Let\'s Encrypt) on the domain so the site is served over HTTPS.'),
        ('3.6', 'Domain configuration', 'Point the production domain (e.g. basecampdigital.com) to the server\'s IP.'),
        ('3.7', 'Environment variables on server', 'All secrets (DB URL, Razorpay keys, JWT secret, SMTP) must be set as environment variables on the server (e.g., via a .env file or PM2 ecosystem config).'),
        ('3.8', 'File storage directory', 'Create a writable directory on the server for uploaded course assets (thumbnails, brochures), and provide the path. Ensure it is publicly accessible via URL or served through Nginx.'),
        ('3.9', 'JWT Secret', 'Generate a strong random string for JWT_SECRET used to sign authentication tokens and share it securely.'),
    ]
)

# ---- PART 2 ----
add_heading(doc, 'Part 2: What We Are Building on the Backend', 1)
doc.add_paragraph('This section explains the backend architecture in plain language so the tech team understands the scope and can plan accordingly.')

add_heading(doc, 'Overview', 2)
doc.add_paragraph(
    'The BCD website is built using Next.js 16 (React framework) with a MySQL database managed through Prisma ORM. '
    'All backend logic runs as secure API routes inside the same Next.js application — there is no separate backend server.'
)

add_heading(doc, 'Database Schema (Key Tables)', 2)
add_table(doc,
    ['Table', 'What It Stores'],
    [
        ('users', 'Registered user accounts — email, hashed password, admin flag, name'),
        ('courses', 'All course metadata — title, slug, price, instructor, thumbnail, publish status'),
        ('sections', 'Ordered sections (chapters) within a course'),
        ('lessons', 'Individual video lessons within a section, linked to Vimeo video IDs'),
        ('enrollments', 'Records of which user is enrolled in which course'),
        ('lesson_progress', 'Tracks which lessons a user has completed'),
        ('course_content', 'Rich structured content per course — learning points, requirements, who it\'s for'),
        ('payment_orders', 'Payment transaction records with Razorpay order/payment IDs and status'),
    ]
)

add_heading(doc, 'Authentication System', 2)
bullets = [
    'Custom JWT-based auth — no third-party auth provider dependency for the main app.',
    'Users sign up or log in via email + password. Passwords are hashed with bcryptjs.',
    'On login, a JWT token is issued (signed with JWT_SECRET) and stored in a secure HTTP-only cookie.',
    'Social login (Google, LinkedIn, Facebook) is handled via OAuth redirects through dedicated API routes.',
    'Admin accounts are not created through any UI — they are manually set in the database (is_admin = true). The admin area at /admin is completely hidden from normal users.',
    'Password reset flow: user requests a reset link via email → receives a time-limited link → sets a new password.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')
doc.add_paragraph()

add_heading(doc, 'API Routes (Backend Endpoints)', 2)
doc.add_paragraph('All routes live under /app/api/ in the Next.js application:')
add_table(doc,
    ['Route Group', 'What It Does'],
    [
        ('/api/auth/login', 'Validates credentials, issues JWT cookie'),
        ('/api/auth/signup', 'Creates a new user account'),
        ('/api/auth/logout', 'Clears the auth cookie'),
        ('/api/auth/me', 'Returns the current logged-in user\'s profile'),
        ('/api/auth/forgot-password', 'Sends a password reset email'),
        ('/api/auth/reset-password', 'Updates the password after reset link is clicked'),
        ('/api/auth/google, /linkedin, /facebook', 'OAuth social login flows'),
        ('/api/enroll', 'Enrolls a user in a course (called after successful payment)'),
        ('/api/progress', 'Marks a lesson as complete / retrieves progress for a course'),
        ('/api/quiz', 'Serves quiz questions and records quiz results'),
        ('/api/certificate', 'Generates a PDF course completion certificate'),
        ('/api/coupons/validate', 'Public endpoint to validate a discount coupon code'),
        ('/api/admin/courses', 'Admin CRUD for courses'),
        ('/api/admin/enrollments', 'Admin view of all enrollments'),
        ('/api/admin/users', 'Admin user management'),
        ('/api/admin/stats', 'Dashboard stats (total users, enrollments, revenue)'),
    ]
)

add_heading(doc, 'Payment Flow', 2)
steps = [
    'User selects a course and clicks "Enroll Now".',
    'Our backend creates a Razorpay payment order (server-side, using the secret key).',
    'The Razorpay checkout modal opens on the frontend with the order ID.',
    'On successful payment, Razorpay calls our webhook (/api/payment-webhook) to confirm.',
    'We verify the payment signature server-side, then create an Enrollment record in the database.',
    'The user is immediately redirected to the course player.',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')
p = doc.add_paragraph()
run = p.add_run('⚠️ Important: Payment confirmation happens via the server-side webhook only — we never trust the frontend to confirm a payment. This requires the Razorpay webhook to be correctly configured and the RAZORPAY_WEBHOOK_SECRET to be set.')
run.bold = True

add_heading(doc, 'Admin Panel', 2)
doc.add_paragraph('A completely hidden admin area at /admin (no public link, no sign-up):')
admin_bullets = [
    'Course Management — Create, edit, delete courses. Upload thumbnails to the server\'s file storage directory. Manage sections and lessons (including assigning Vimeo video IDs).',
    'Coupon Management — Create discount coupons with percentage off, usage limits, and expiry dates.',
    'Enrollment Management — View all enrollments. Manually enroll users (e.g., for corporate/offline purchases).',
    'User Management — View registered users, promote users to admin.',
    'Dashboard — Quick stats: total courses, enrollments, registered users.',
]
for b in admin_bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading(doc, 'Content Delivery', 2)
content_bullets = [
    'Video lessons are hosted on Vimeo and embedded via the @vimeo/player library. Videos are domain-restricted — they can only be played on the BCD website.',
    'Course thumbnails and brochures are uploaded to and served from a directory on the team\'s own server, accessible via a public URL path configured through Nginx.',
    'Dynamic course content (learning points, requirements, "who this is for") is stored in the MySQL database and fetched at page load — nothing is hardcoded in the frontend.',
]
for b in content_bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading(doc, 'Key Technology Stack Summary', 2)
add_table(doc,
    ['Layer', 'Technology'],
    [
        ('Frontend + Backend', 'Next.js 16 (React 19)'),
        ('Language', 'TypeScript'),
        ('Database', 'MySQL on SiteGround (via Prisma ORM)'),
        ('Auth', 'Custom JWT + bcryptjs'),
        ('Payments', 'Razorpay'),
        ('Video', 'Vimeo'),
        ('File Storage', 'Team\'s own server (Nginx-served directory)'),
        ('Email', 'Nodemailer (SMTP) — already configured'),
        ('Deployment', 'Team\'s own server (PM2 + Nginx)'),
        ('Analytics', 'Vercel Analytics'),
    ]
)

doc.add_paragraph()
p = doc.add_paragraph('End of Document')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

output_path = '/Users/parthlodat/Desktop/BCD_Backend_Requirements.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
